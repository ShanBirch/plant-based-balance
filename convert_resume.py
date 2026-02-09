import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_resume_docx(html_path, output_path):
    if not os.path.exists(html_path):
        print(f"Error: File not found at {html_path}")
        return

    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')

    doc = Document()

    # Style modifications (optional, keeping defaults usually looks okay or we can tweak)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Header ---
    header_tag = soup.find('header')
    if header_tag:
        # Name (H1)
        h1 = header_tag.find('h1')
        if h1:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(h1.get_text().strip())
            run.bold = True
            run.font.size = Pt(24)
            run.font.color.rgb = None # Default or set color

        # Role
        role = header_tag.find(class_='role-title')
        if role:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(role.get_text().strip())
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = None # We could try to match standard blue like RGB(41, 128, 185)

        # Contact Info
        contact_div = header_tag.find(class_='contact-info')
        if contact_div:
            spans = contact_div.find_all('span')
            contact_text = " | ".join([s.get_text().strip() for s in spans])
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.add_run(contact_text)

        doc.add_paragraph() # Spacer

    # --- Sections ---
    sections = soup.find_all('section')
    for section in sections:
        # Section Title (H2)
        h2 = section.find('h2')
        if h2:
            h = doc.add_heading(h2.get_text().strip(), level=1)
            # Optional: Add bottom border logic or just rely on styles

        # Process children elements
        # We iterate over children to maintain order
        for child in section.children:
            if child.name == 'p':
                text = child.get_text().strip()
                if text:
                    doc.add_paragraph(text)
            
            elif child.name == 'ul':
                for li in child.find_all('li'):
                    doc.add_paragraph(li.get_text().strip(), style='List Bullet')

            elif child.name == 'div' and 'experience-item' in child.get('class', []):
                # Job Title & Date
                main_line = child.find(class_='job-main-line')
                title_text = ""
                date_text = ""
                if main_line:
                    jt = main_line.find(class_='job-title')
                    if jt: title_text = jt.get_text().strip()
                    dr = main_line.find(class_='date-range')
                    if dr: date_text = dr.get_text().strip()
                
                # We can put Title and Date on one line spread out, or just formatted
                p = doc.add_paragraph()
                if title_text:
                    r = p.add_run(title_text)
                    r.bold = True
                    r.font.size = Pt(12)
                if date_text:
                    p.add_run(f"  ({date_text})").italic = True

                # Company & Location
                comp_name = child.find(class_='company-name')
                if comp_name:
                    p = doc.add_paragraph()
                    # Location often inside, let's extract text directly
                    p.add_run(comp_name.get_text().strip()).bold = True

                # Description P
                desc_p = child.find('p')
                if desc_p:
                    doc.add_paragraph(desc_p.get_text().strip())

                # Description UL
                desc_ul = child.find('ul')
                if desc_ul:
                    for li in desc_ul.find_all('li'):
                        # Handle internal formatting like strong tags in LI?
                        # Simple get_text for now
                        p = doc.add_paragraph(style='List Bullet')
                        # Check for bold starts
                        if li.find('strong'):
                            # Elementary parsing: if starts with strong, bold it
                            # better: iterate contents
                            for content in li.contents:
                                if content.name == 'strong':
                                    p.add_run(content.get_text()).bold = True
                                elif isinstance(content, str):
                                    p.add_run(content)
                        else:
                            p.add_run(li.get_text().strip())
                
                doc.add_paragraph() # Spacer after item

            elif child.name == 'div' and 'skills-grid' in child.get('class', []):
                # Skills Grid logic
                # Check if we have div columns (ML resume) or just ULs (Standard resume)
                cols = child.find_all('div', recursive=False)
                if cols:
                    for col in cols:
                        cat = col.find(class_='skills-category')
                        if cat:
                            p = doc.add_paragraph()
                            p.add_run(cat.get_text().strip()).bold = True
                        
                        ul = col.find('ul')
                        if ul:
                            for li in ul.find_all('li'):
                                 doc.add_paragraph(li.get_text().strip(), style='List Bullet')
                        doc.add_paragraph() # Spacer
                else:
                    # Direct ULs case
                    uls = child.find_all('ul', recursive=False)
                    for ul in uls:
                        for li in ul.find_all('li'):
                            doc.add_paragraph(li.get_text().strip(), style='List Bullet')
                        doc.add_paragraph() # Spacer

    doc.save(output_path)
    print(f"Successfully created {output_path}")

if __name__ == "__main__":
    html_file = "C:/Users/shann/.gemini/antigravity/plant_based_balance/resume.html"
    docx_file = "C:/Users/shann/.gemini/antigravity/plant_based_balance/shannonbirch_resume.docx"
    create_resume_docx(html_file, docx_file)
